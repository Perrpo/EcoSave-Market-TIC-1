"""
Script para generar un documento Word (.docx) con el contenido
del póster/presentación de EcoSave Market.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def set_cell_shading(cell, color_hex):
    """Aplica color de fondo a una celda."""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elem)

def add_styled_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        if level == 1:
            run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)  # green-600
            run.font.size = Pt(22)
        elif level == 2:
            run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)  # green-500
            run.font.size = Pt(16)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)  # slate-700
            run.font.size = Pt(13)
    return heading

def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x33, 0x41, 0x55)

def add_normal_paragraph(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 80)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)

def main():
    doc = Document()
    
    # ── Configurar márgenes ──
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Configurar estilo Normal ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    # ═══════════════════════════════════════════════════════
    # PORTADA
    # ═══════════════════════════════════════════════════════
    
    # Espacio superior
    for _ in range(4):
        doc.add_paragraph()
    
    # Badge UPB
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Proyecto Académico · Universidad Pontificia Bolivariana')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x10, 0xB9, 0x81)
    run.bold = True

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('EcoSave Market')
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
    run.bold = True
    
    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Plataforma Anti-Desperdicio de Alimentos')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Descripción
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run(
        'Sistema de gestión multi-rol para productos próximos a vencer en minimercados '
        'y tiendas de barrio de Medellín, conectándolos con ONGs y consumidores para '
        'reducir el desperdicio y promover el aprovechamiento sostenible.'
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Palabras clave
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    keywords = ['🍎 Alimentos próximos a vencer', '♻️ Economía circular', 
                '🌱 Aprovechamiento sostenible', '📉 Reducción de desperdicio', 
                '🤝 Conexión social']
    run = p.add_run('  |  '.join(keywords))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # Info institucional
    for _ in range(4):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Universidad Pontificia Bolivariana')
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = RGBColor(0xC4, 0x12, 0x30)  # UPB red
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Curso TIC 1 · Facultad de Ingeniería · Medellín, Colombia · 2026')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # ── Salto de página ──
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 01 — PROBLEMA
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '01 — Problema', level=1)
    add_styled_heading(doc, 'Una crisis alimentaria evitable', level=2)
    add_normal_paragraph(doc, 
        'Colombia enfrenta una paradoja: millones de toneladas de alimentos se desperdician '
        'mientras millones de personas pasan hambre.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # Tabla de problemas
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Encabezados
    headers = ['Indicador', 'Descripción', 'Dato Clave']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Fila 1
    problems = [
        ['📊 Desperdicio masivo de alimentos',
         'En Colombia se desperdician cerca de 9.76 millones de toneladas de alimentos al año, '
         'equivalentes al 34% de la producción total. El 20.6% se pierde en distribución y retail.',
         '9.76M toneladas/año desperdiciadas'],
        ['⚠️ Inseguridad alimentaria',
         'Aproximadamente 14.4 millones de colombianos (27.6% de la población) se encuentran '
         'en situación de inseguridad alimentaria moderada o grave.',
         '14.4M personas en inseguridad alimentaria'],
        ['🏪 Desconexión digital en tiendas',
         'Las tiendas de barrio y minimercados carecen de herramientas digitales para gestionar '
         'inventarios y canalizar excedentes. Las ONGs operan con procesos desarticulados sin '
         'coordinación eficiente.',
         '0 plataformas integradas para microcomercio']
    ]
    
    for i, row_data in enumerate(problems):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 02 — SOLUCIÓN PROPUESTA
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '02 — Solución Propuesta', level=1)
    add_styled_heading(doc, 'Intermediario digital para salvar alimentos', level=2)
    add_normal_paragraph(doc, 
        'Una plataforma web que conecta establecimientos, ONGs y consumidores de forma '
        'eficiente, trazable y transparente.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    add_styled_heading(doc, '¿Qué construimos?', level=3)
    add_normal_paragraph(doc,
        'EcoSave Market es una plataforma web multi-rol que funciona como intermediario digital '
        'entre minimercados, tiendas de barrio y pequeños negocios alimentarios con ONGs, bancos '
        'de alimentos y consumidores. Las tiendas registran productos próximos a vencer '
        '(especificando cantidades, fechas y condiciones), mientras que las ONGs y consumidores '
        'consultan disponibilidad y coordinan su retiro. Todo con trazabilidad completa y '
        'generación automática de certificados de donación alineados con la Ley 2380 de 2024 '
        'para beneficios tributarios del 37%.')

    # Funcionalidades
    add_styled_heading(doc, 'Funcionalidades principales', level=3)
    features = [
        'Gestión de excedentes — Registro y publicación de productos próximos a vencer con estados visuales tipo semáforo.',
        'Redistribución inteligente — Algoritmo de asignación basado en prioridad, distancia geográfica y capacidad de las ONGs.',
        'Certificados digitales — Generación automática de certificados PDF con cálculo de deducción tributaria del 37%.',
        'Dashboard por rol — Paneles diferenciados con métricas de impacto social y ambiental en tiempo real.',
        'Mapa interactivo — Visualización de ONGs cercanas, puntos de recolección y rutas de entrega optimizadas.',
        'Notificaciones y trazabilidad — Alertas de vencimiento, confirmación de donaciones y trazabilidad completa del flujo.'
    ]
    add_bullet_list(doc, features)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 03 — USUARIOS Y ROLES
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '03 — Usuarios y Roles', level=1)
    add_styled_heading(doc, 'Tres actores, un ecosistema', level=2)
    add_normal_paragraph(doc, 
        'Cada usuario tiene necesidades específicas que la plataforma resuelve de forma personalizada.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # -- Tienda / Minimercado --
    add_styled_heading(doc, '🏪 Tienda / Minimercado (Rol: Supermarket)', level=3)
    add_normal_paragraph(doc,
        'Pequeños negocios alimentarios que necesitan reducir pérdidas por productos próximos a '
        'vencer y mejorar su responsabilidad social.')
    tienda_needs = [
        'Registrar productos próximos a vencer con fecha y unidades',
        'Publicar donaciones disponibles para ONGs',
        'Generar certificados PDF de donación (Ley 2380)',
        'Visualizar métricas de impacto y ahorro tributario',
        'Recibir notificaciones de solicitudes de ONGs',
        'Enviar reportes consolidados por email'
    ]
    add_bullet_list(doc, tienda_needs)

    # -- ONG / Fundación --
    add_styled_heading(doc, '🤝 ONG / Fundación (Rol: ONG)', level=3)
    add_normal_paragraph(doc,
        'Organizaciones que redistribuyen alimentos a poblaciones vulnerables y necesitan acceso '
        'organizado y ágil a excedentes.')
    ong_needs = [
        'Explorar donaciones disponibles en tiempo real',
        'Solicitar alimentos (parcial o total)',
        'Confirmar recepción y cerrar el ciclo',
        'Consultar mapa de puntos de recolección',
        'Acceder al historial de recepciones',
        'Visualizar estadísticas de impacto social'
    ]
    add_bullet_list(doc, ong_needs)

    # -- Consumidor --
    add_styled_heading(doc, '🧑‍🤝‍🧑 Consumidor (Rol: Consumer — proyectado)', level=3)
    add_normal_paragraph(doc,
        'Personas que buscan alimentos de calidad a precios reducidos, promoviendo el consumo '
        'responsable y la economía circular.')
    consumer_needs = [
        'Acceder a productos a precios reducidos',
        'Consumo responsable y sostenible',
        'Visualizar impacto ambiental de sus compras',
        'Calificar experiencias de adquisición',
        'Marcar establecimientos como favoritos',
        'Recibir alertas de productos disponibles'
    ]
    add_bullet_list(doc, consumer_needs)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 04 — RESULTADO / VALOR
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '04 — Resultado / Valor', level=1)
    add_styled_heading(doc, 'Impacto medible y verificable', level=2)
    add_normal_paragraph(doc,
        'Funcionalidades implementadas que generan valor económico, social y ambiental.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # Estadísticas clave
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    stats = [
        ('12', 'Historias de usuario implementadas'),
        ('37%', 'Deducción tributaria (Ley 2380)'),
        ('3', 'Roles con dashboards independientes'),
        ('100%', 'Trazabilidad del flujo de donación')
    ]
    
    for i, (value, label) in enumerate(stats):
        cell_val = table.rows[0].cells[i]
        cell_val.text = value
        for paragraph in cell_val.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x05, 0x96, 0x69)
        
        cell_label = table.rows[1].cells[i]
        cell_label.text = label
        for paragraph in cell_label.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(9)

    doc.add_paragraph()  # Espaciado

    # Resultados detallados
    add_styled_heading(doc, 'Resultados de la implementación', level=3)
    results = [
        '💰 Reducción de pérdidas económicas — Las tiendas convierten pérdidas en oportunidades mediante donaciones con beneficios tributarios y ventas a precio reducido.',
        '🌍 Impacto ambiental positivo — Reducción de emisiones de GEI al evitar que alimentos aptos terminen en vertederos (8-10% de emisiones globales).',
        '📊 Dashboard con métricas de impacto — Productos donados, kg rescatados, ONGs activas y estadísticas diferenciadas por rol en tiempo real.',
        '📄 Certificados automáticos de donación — Generación de PDF con cálculo de deducción tributaria del 37%, alineado con la Ley 2380 de 2024.',
        '🔄 Redistribución inteligente de sobrantes — Algoritmo que redistribuye automáticamente excedentes no reclamados a otras ONGs disponibles.',
        '🗺️ Geolocalización y logística — Mapa interactivo con ONGs cercanas, puntos de recolección y filtros por especialidad para optimizar entregas.'
    ]
    add_bullet_list(doc, results)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 05 — DEMOSTRACIÓN VISUAL
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '05 — Demostración Visual', level=1)
    add_styled_heading(doc, 'La plataforma en acción', level=2)
    add_normal_paragraph(doc,
        'Capturas del sistema, dashboard y flujo de proceso implementado.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # Insertar imágenes si existen
    docs_dir = os.path.dirname(os.path.abspath(__file__))
    
    dashboard_img = os.path.join(docs_dir, 'dashboard_screenshot.png')
    if os.path.exists(dashboard_img):
        add_styled_heading(doc, 'Dashboard del Supermercado', level=3)
        add_normal_paragraph(doc,
            'Panel principal con métricas de impacto, productos próximos a vencer y acciones rápidas de donación.',
            size=10, color=RGBColor(0x64, 0x74, 0x8B))
        doc.add_picture(dashboard_img, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    process_img = os.path.join(docs_dir, 'process_flow.png')
    if os.path.exists(process_img):
        add_styled_heading(doc, 'Flujo del Proceso', level=3)
        add_normal_paragraph(doc,
            'Desde el registro de productos hasta la generación del certificado de donación.',
            size=10, color=RGBColor(0x64, 0x74, 0x8B))
        doc.add_picture(process_img, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Sistema semáforo
    add_styled_heading(doc, 'Sistema de estados visuales tipo semáforo', level=3)
    add_normal_paragraph(doc,
        'Indicadores de color según proximidad de vencimiento del producto:')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    semaforo_headers = ['Estado', 'Color', 'Condición']
    for i, h in enumerate(semaforo_headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    
    semaforo_data = [
        ['🟢 Seguro', 'Verde', '> 7 días para vencer'],
        ['🟡 Próximo', 'Amarillo', '3–7 días para vencer'],
        ['🔴 Urgente', 'Rojo', '< 3 días para vencer']
    ]
    for i, row_data in enumerate(semaforo_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for p in table.rows[i + 1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 06 — STACK TECNOLÓGICO
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '06 — Stack Tecnológico', level=1)
    add_styled_heading(doc, 'Arquitectura robusta y escalable', level=2)
    add_normal_paragraph(doc,
        'Tecnologías modernas organizadas en capas para máxima mantenibilidad.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # Diagrama de arquitectura como tabla
    add_styled_heading(doc, 'Arquitectura por capas', level=3)
    
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    table.rows[0].cells[0].text = 'Capa'
    table.rows[0].cells[1].text = 'Tecnologías'
    for p in table.rows[0].cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    for p in table.rows[0].cells[1].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)
    
    arch_data = [
        ['Frontend', 'React 19 · TypeScript 5.9 · Vite 7 · Vanilla CSS · React Router 7 · Framer Motion'],
        ['Backend', 'AdonisJS 6 · TypeScript 5.8 · PDFKit · Nodemailer · JWT Auth'],
        ['Base de Datos', 'Supabase · PostgreSQL · Row Level Security'],
        ['Infraestructura', 'Docker · Docker Compose · Nginx']
    ]
    
    for i, (layer, techs) in enumerate(arch_data):
        table.rows[i + 1].cells[0].text = layer
        table.rows[i + 1].cells[1].text = techs
        for p in table.rows[i + 1].cells[0].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        for p in table.rows[i + 1].cells[1].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

    doc.add_paragraph()

    # Tech stack detallado
    add_styled_heading(doc, 'Detalle del stack', level=3)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    tech_headers = ['Tecnología', 'Versión', 'Capa', 'Propósito']
    for i, h in enumerate(tech_headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    
    tech_data = [
        ['⚛️ React', 'v19', 'Frontend', 'Interfaces reactivas SPA'],
        ['📘 TypeScript', 'v5.9', 'Frontend', 'Tipado estático seguro'],
        ['⚡ Vite', 'v7', 'Frontend', 'Build tool y HMR'],
        ['🟣 AdonisJS', 'v6', 'Backend', 'API REST con capas'],
        ['🟢 Supabase', 'PostgreSQL', 'BD', 'BaaS con RLS'],
        ['🐳 Docker', 'Compose', 'Infra', 'Despliegue reproducible']
    ]
    
    for i, row_data in enumerate(tech_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for p in table.rows[i + 1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 07 — SIGUIENTES PASOS
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '07 — Siguientes Pasos', level=1)
    add_styled_heading(doc, 'Roadmap de evolución', level=2)
    add_normal_paragraph(doc,
        '¿Qué falta para llevar el proyecto a una implementación real?',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # Fase 1
    add_styled_heading(doc, 'Fase 1 — Corto Plazo', level=3)
    fase1 = [
        'Diseño visual mejorado y UX optimizada — Interfaz más moderna, accesible y responsive. Adaptación completa a dispositivos móviles y optimización de la navegación para todos los perfiles.',
        'Sistema de estados visuales tipo semáforo — Indicadores de color verde/amarillo/rojo para la proximidad de vencimiento. Identificación rápida para decisiones eficientes en el dashboard.'
    ]
    add_bullet_list(doc, fase1)

    # Fase 2
    add_styled_heading(doc, 'Fase 2 — Mediano Plazo', level=3)
    fase2 = [
        'Reservas y límites por donación — Sistema de reservas temporales y límites de cantidad por usuario para evitar sobreasignación y garantizar distribución equitativa.',
        'Sistema de reputación y favoritos — Calificaciones de experiencia para generar confianza, y marcado de organizaciones/establecimientos favoritos para interacciones recurrentes.'
    ]
    add_bullet_list(doc, fase2)

    # Fase 3
    add_styled_heading(doc, 'Fase 3 — Largo Plazo', level=3)
    fase3 = [
        'Integración con sistemas de inventario — Conectores API con Siigo, Alegra y Google Sheets para sincronización automática de productos. Webhooks para actualizaciones en tiempo real.',
        'Escalabilidad a supermercados y cadenas — Expansión progresiva hacia cadenas regionales de mayor tamaño, manteniendo la arquitectura modular y escalable del sistema.'
    ]
    add_bullet_list(doc, fase3)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 08 — PLAN DE NEGOCIO
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '08 — Plan de Negocio', level=1)
    add_styled_heading(doc, 'Modelo de valor sostenible', level=2)
    add_normal_paragraph(doc,
        'Propuesta de valor, cliente objetivo y generación de valor para cada actor del ecosistema.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    # Propuesta de Valor
    add_styled_heading(doc, '💎 Propuesta de Valor', level=3)
    add_normal_paragraph(doc,
        'Plataforma digital que transforma el desperdicio en oportunidad, conectando oferta y '
        'demanda de alimentos próximos a vencer con trazabilidad completa, beneficios tributarios '
        'automatizados y métricas de impacto social.')
    propuesta = [
        'Reducción de pérdidas económicas por productos vencidos',
        'Beneficios tributarios del 37% por donaciones (Ley 2380)',
        'Trazabilidad y certificación digital de donaciones',
        'Impacto medible en reducción de desperdicio'
    ]
    add_bullet_list(doc, propuesta)

    # Cliente Objetivo
    add_styled_heading(doc, '🎯 Cliente Objetivo', level=3)
    add_normal_paragraph(doc,
        'Inicialmente enfocado en el ecosistema de Medellín, con escalabilidad hacia otras regiones de Colombia.')
    clientes = [
        'Primario: Minimercados y tiendas de barrio de Medellín',
        'Secundario: ONGs, fundaciones y bancos de alimentos',
        'Terciario: Consumidores conscientes y responsables',
        'Futuro: Supermercados y cadenas regionales'
    ]
    add_bullet_list(doc, clientes)

    # Cómo Genera Valor
    add_styled_heading(doc, '⚙️ Cómo Genera Valor', level=3)
    add_normal_paragraph(doc,
        'Un ecosistema donde todos los actores obtienen beneficios tangibles y medibles.')
    valor = [
        'Para tiendas: Menos pérdidas, más ahorro fiscal, mejor imagen',
        'Para ONGs: Acceso organizado a alimentos, logística optimizada',
        'Para consumidores: Productos de calidad a menor precio',
        'Para la sociedad: Reducción de inseguridad alimentaria y emisiones'
    ]
    add_bullet_list(doc, valor)

    # Marco Legal
    add_styled_heading(doc, '📋 Marco Legal Soportante', level=3)
    add_normal_paragraph(doc,
        'Respaldo normativo sólido que incentiva y facilita la operación de la plataforma.')
    marco = [
        'Ley 1990/2019: Política contra pérdida y desperdicio',
        'Ley 2380/2024: Donación de alimentos y beneficios tributarios',
        'Ley 1581/2012: Protección de datos personales',
        'ODS 2 y 12: Hambre Cero y Consumo Responsable'
    ]
    add_bullet_list(doc, marco)

    add_separator(doc)

    # ═══════════════════════════════════════════════════════
    # SECCIÓN 09 — EQUIPO
    # ═══════════════════════════════════════════════════════
    add_styled_heading(doc, '09 — Equipo', level=1)
    add_styled_heading(doc, 'El equipo detrás de EcoSave', level=2)
    add_normal_paragraph(doc,
        'Proyecto académico de transformación digital — Universidad Pontificia Bolivariana, Medellín.',
        italic=True, color=RGBColor(0x64, 0x74, 0x8B))

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    team_headers = ['Integrante', 'Rol', 'Responsabilidades']
    for i, h in enumerate(team_headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
    
    team_data = [
        ['Integrante A', 'Frontend Developer', 'React · TypeScript · Diseño UI/UX · Componentes reutilizables'],
        ['Integrante B', 'Backend Developer', 'AdonisJS · API REST · Lógica de negocio · Autenticación'],
        ['Integrante C', 'Base de Datos & DevOps', 'Supabase · PostgreSQL · Docker · Despliegue'],
        ['Integrante D', 'Investigación & QA', 'Documentación · Testing · Estado del arte · Marco legal']
    ]
    
    for i, row_data in enumerate(team_data):
        for j, val in enumerate(row_data):
            table.rows[i + 1].cells[j].text = val
            for p in table.rows[i + 1].cells[j].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    # ── PIE DE PÁGINA ──
    doc.add_paragraph()
    add_separator(doc)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('EcoSave Market — Plataforma Anti-Desperdicio de Alimentos')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    run.italic = True
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Economía Circular · Desarrollo Sostenible · UPB 2026')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    # ═══════════════════════════════════════════════════════
    # GUARDAR
    # ═══════════════════════════════════════════════════════
    output_path = os.path.join(docs_dir, 'EcoSave_Market_Presentacion.docx')
    doc.save(output_path)
    print(f'[OK] Documento generado exitosamente: {output_path}')


if __name__ == '__main__':
    main()
